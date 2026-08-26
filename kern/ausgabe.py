# -*- coding: utf-8 -*-
"""
Schreibt das fertige Transkript als PDF, Word-Datei oder Textdatei.
Alle Dateien landen im Ordner 'ergebnisse'.

Jeder Absatz kann einer Person zugeordnet sein. Wer mitkommt und wer
nicht, entscheidet der Aufrufer ueber die schon gefilterte Absatzliste.
"""

import re
from datetime import datetime
from pathlib import Path

from kern.absaetze import name_von, personen_liste, zeitstempel

ERGEBNIS_ORDNER = Path(__file__).resolve().parent.parent / "ergebnisse"

# Farben fuer die Personen, wiederholen sich ab der neunten Person.
FARBEN = ["#2563eb", "#dc2626", "#16a34a", "#d97706",
          "#7c3aed", "#0891b2", "#db2777", "#65a30d"]


def farbe_von(person):
    person = int(person or 0)
    if person <= 0:
        return "#6b7280"
    return FARBEN[(person - 1) % len(FARBEN)]


def _sauberer_name(text):
    text = (text or "Transkript").strip()
    text = re.sub(r"[^\w\s\-\.]", "", text, flags=re.UNICODE)
    text = re.sub(r"\s+", "_", text).strip("._")
    return text[:60] or "Transkript"


def _zielpfad(titel, endung):
    ERGEBNIS_ORDNER.mkdir(parents=True, exist_ok=True)
    stempel = datetime.now().strftime("%Y-%m-%d_%H-%M")
    pfad = ERGEBNIS_ORDNER / ("%s_%s.%s" % (stempel, _sauberer_name(titel), endung))

    zaehler = 2
    while pfad.exists():
        pfad = ERGEBNIS_ORDNER / (
            "%s_%s_%d.%s" % (stempel, _sauberer_name(titel), zaehler, endung))
        zaehler += 1
    return pfad


def _kopfzeilen(absaetze, orion_an, dauer_sekunden, namen=None):
    woerter = sum(len(a["text"].split()) for a in absaetze
                  if not a.get("geraeusch"))
    leute = [e for e in personen_liste(absaetze) if e["person"] > 0]

    zeilen = [
        ("Erstellt", datetime.now().strftime("%d.%m.%Y um %H:%M Uhr")),
        ("Laenge", zeitstempel(dauer_sekunden) if dauer_sekunden else "unbekannt"),
        ("Woerter", str(woerter)),
        ("Absaetze", str(len(absaetze))),
    ]
    if leute:
        zeilen.append(("Personen", ", ".join(
            name_von(e["person"], namen) for e in leute)))
    zeilen.append(("Orion-Funktion",
                   "eingeschaltet" if orion_an else "ausgeschaltet"))
    return zeilen


# ----------------------------------------------------------------------
def als_txt(titel, absaetze, orion_an=True, mit_zeit=False, dauer_sekunden=0,
            namen=None, mit_person=True):
    pfad = _zielpfad(titel, "txt")

    zeilen = [titel, "=" * len(titel), ""]
    for schild, wert in _kopfzeilen(absaetze, orion_an, dauer_sekunden, namen):
        zeilen.append("%-16s %s" % (schild + ":", wert))
    zeilen += ["", "-" * 60, ""]

    for a in absaetze:
        kopf = []
        if mit_zeit:
            kopf.append("[%s]" % zeitstempel(a["start"]))
        if mit_person and not a.get("geraeusch"):
            kopf.append("%s:" % name_von(a.get("person"), namen))
        if kopf:
            zeilen.append(" ".join(kopf))
        zeilen.append(a["text"])
        zeilen.append("")

    pfad.write_text("\n".join(zeilen), encoding="utf-8")
    return pfad


# ----------------------------------------------------------------------
def als_docx(titel, absaetze, orion_an=True, mit_zeit=False, dauer_sekunden=0,
             namen=None, mit_person=True):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        raise RuntimeError("python-docx fehlt. Bitte installieren.bat ausfuehren.")

    pfad = _zielpfad(titel, "docx")
    dok = Document()

    normal = dok.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    dok.add_heading(titel, level=0)

    for schild, wert in _kopfzeilen(absaetze, orion_an, dauer_sekunden, namen):
        p = dok.add_paragraph()
        lauf = p.add_run("%s: " % schild)
        lauf.bold = True
        lauf.font.size = Pt(9)
        p.add_run(wert).font.size = Pt(9)
        p.paragraph_format.space_after = Pt(0)

    dok.add_paragraph()

    for a in absaetze:
        kopf = dok.add_paragraph()
        kopf.paragraph_format.space_after = Pt(1)
        etwas_im_kopf = False

        if mit_person and not a.get("geraeusch"):
            lauf = kopf.add_run(name_von(a.get("person"), namen))
            lauf.bold = True
            lauf.font.size = Pt(9.5)
            hex_farbe = farbe_von(a.get("person")).lstrip("#")
            lauf.font.color.rgb = RGBColor(
                int(hex_farbe[0:2], 16), int(hex_farbe[2:4], 16),
                int(hex_farbe[4:6], 16))
            etwas_im_kopf = True

        if mit_zeit:
            lauf = kopf.add_run(("   " if etwas_im_kopf else "")
                                + zeitstempel(a["start"]))
            lauf.font.size = Pt(8.5)
            lauf.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            etwas_im_kopf = True

        if not etwas_im_kopf:
            kopf._element.getparent().remove(kopf._element)

        p = dok.add_paragraph(a["text"])
        p.paragraph_format.space_after = Pt(10)
        if a.get("geraeusch"):
            for lauf in p.runs:
                lauf.italic = True
                lauf.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    dok.save(str(pfad))
    return pfad


# ----------------------------------------------------------------------
def als_pdf(titel, absaetze, orion_an=True, mit_zeit=False, dauer_sekunden=0,
            namen=None, mit_person=True):
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.lib import colors
        from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                        HRFlowable, KeepTogether)
    except ImportError:
        raise RuntimeError("reportlab fehlt. Bitte installieren.bat ausfuehren.")

    from xml.sax.saxutils import escape

    pfad = _zielpfad(titel, "pdf")

    dok = SimpleDocTemplate(
        str(pfad), pagesize=A4,
        leftMargin=22 * mm, rightMargin=22 * mm,
        topMargin=20 * mm, bottomMargin=18 * mm,
        title=titel, author="Transkript",
    )

    basis = getSampleStyleSheet()

    stil_titel = ParagraphStyle(
        "TitelGross", parent=basis["Title"], fontName="Helvetica-Bold",
        fontSize=19, leading=23, spaceAfter=4, alignment=0,
        textColor=colors.HexColor("#111111"))
    stil_kopf = ParagraphStyle(
        "Kopf", parent=basis["Normal"], fontName="Helvetica",
        fontSize=8.5, leading=12, textColor=colors.HexColor("#666666"))
    stil_sprecher = ParagraphStyle(
        "Sprecher", parent=basis["Normal"], fontName="Helvetica-Bold",
        fontSize=9.5, leading=12, spaceAfter=1)
    stil_text = ParagraphStyle(
        "Fliesstext", parent=basis["Normal"], fontName="Helvetica",
        fontSize=10.5, leading=15.5, spaceAfter=9,
        textColor=colors.HexColor("#1a1a1a"))
    stil_geraeusch = ParagraphStyle(
        "Geraeusch", parent=stil_text, fontName="Helvetica-Oblique",
        textColor=colors.HexColor("#777777"))

    inhalt = [Paragraph(escape(titel), stil_titel)]

    kopf = "&nbsp;&nbsp;|&nbsp;&nbsp;".join(
        "%s: %s" % (escape(s), escape(w))
        for s, w in _kopfzeilen(absaetze, orion_an, dauer_sekunden, namen))
    inhalt.append(Paragraph(kopf, stil_kopf))
    inhalt.append(Spacer(1, 5))
    inhalt.append(HRFlowable(width="100%", thickness=0.6,
                             color=colors.HexColor("#dddddd")))
    inhalt.append(Spacer(1, 9))

    for a in absaetze:
        block = []
        teile = []
        if mit_person and not a.get("geraeusch"):
            teile.append('<font color="%s">%s</font>'
                         % (farbe_von(a.get("person")),
                            escape(name_von(a.get("person"), namen))))
        if mit_zeit:
            teile.append('<font size="8" color="#999999">%s</font>'
                         % zeitstempel(a["start"]))
        if teile:
            block.append(Paragraph("&nbsp;&nbsp;".join(teile), stil_sprecher))

        block.append(Paragraph(
            escape(a["text"]),
            stil_geraeusch if a.get("geraeusch") else stil_text))

        inhalt.append(KeepTogether(block) if len(block) > 1 else block[0])

    def seitenzahl(leinwand, dokument):
        leinwand.saveState()
        leinwand.setFont("Helvetica", 8)
        leinwand.setFillColor(colors.HexColor("#999999"))
        leinwand.drawRightString(A4[0] - 22 * mm, 12 * mm,
                                 "Seite %d" % dokument.page)
        leinwand.restoreState()

    dok.build(inhalt, onFirstPage=seitenzahl, onLaterPages=seitenzahl)
    return pfad


# ----------------------------------------------------------------------
SCHREIBER = {"pdf": als_pdf, "docx": als_docx, "txt": als_txt}


def schreiben(format_name, titel, absaetze, orion_an=True, mit_zeit=False,
              dauer_sekunden=0, namen=None, mit_person=True):
    if format_name not in SCHREIBER:
        raise ValueError("Unbekanntes Format: %s" % format_name)
    if not absaetze:
        raise ValueError("Es gibt nichts zu speichern, die Auswahl ist leer.")
    return SCHREIBER[format_name](
        titel, absaetze, orion_an=orion_an, mit_zeit=mit_zeit,
        dauer_sekunden=dauer_sekunden, namen=namen, mit_person=mit_person)
